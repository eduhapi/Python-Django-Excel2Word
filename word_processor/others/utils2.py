import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def extract_tables_from_excel(excel_file_path):
    wb = openpyxl.load_workbook(excel_file_path)
    extracted_tables = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        table_data = []
        for row in ws.iter_rows(values_only=True):
            table_data.append(row)
        extracted_tables.append((sheet, table_data))
    return extracted_tables

def create_word_document(tables, output_path):
    doc = Document()
    
    # Group officers by status
    grouped_tables = {'Active': [], 'Suspended': [], 'Retired': []}
    for sheet_name, table_data in tables:
        for row in table_data[1:]:  # Skip header row
            id_, username, pno, status = row[:4]  # Assuming columns are in order
            officer_info = f"ID: {id_}, Username: {username}, P/No: {pno}"
            if status in ['active', 'suspended', 'retired']:
               grouped_tables.setdefault(status.capitalize(), []).append(officer_info)
        else:
            # Handle unexpected status values (optional)
            print(f"Warning: Unexpected status '{status}' encountered.")
          #  grouped_tables[status].append(officer_info)

    # Create tables for each status
    for status, officer_info_list in grouped_tables.items():
        if officer_info_list:
            doc.add_heading(status, level=1)
            table = doc.add_table(rows=len(officer_info_list), cols=1)
            for i, officer_info in enumerate(officer_info_list):
                cell = table.cell(i, 0)
                cell.text = officer_info
    doc.save(output_path)
