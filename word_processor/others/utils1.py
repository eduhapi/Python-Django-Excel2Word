import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_tables_from_excel(excel_file_path):
    wb = openpyxl.load_workbook(excel_file_path)
    extracted_tables = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        table_data = []
        for row in ws.iter_rows(values_only=True):
            table_data.append(row)
        extracted_tables.append((sheet, table_data))
    return extracted_tables

def create_word_document(tables):
    doc = Document()
    for sheet_name, table_data in tables:
        doc.add_heading(sheet_name, level=1)
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = str(cell)
                table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.cell(i, j).vertical_anchor = 'middle'
                table.cell(i, j).paragraphs[0].runs[0].font.size = Pt(10)
    return doc
