import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def extract_tables_from_excel(excel_file_path):
    wb = openpyxl.load_workbook(excel_file_path)
    extracted_tables = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        table_data = []
        for row in ws.iter_rows(values_only=True):
            # Extract only the columns of interest: OS Ticket Number, Script Mapping, UAT Status
            extracted_row = (row[0], row[2], row[6])  # Assuming columns indices start from 0
            table_data.append(extracted_row)
        extracted_tables.append((sheet, table_data))
    return extracted_tables

def create_word_document(tables, output_path):
    doc = Document()
    
    # Group data by status
    grouped_tables = {'Pass': [], 'Fail': [], 'Not Executed': [], 'Not Applicable': []}
    for sheet_name, table_data in tables:
        for row in table_data[1:]:  # Skip header row
            os_ticket, script_mapping, uat_status = row[:3]  # Extracted columns
            if uat_status in ['Pass', 'Fail', 'Not Executed', 'Not Applicable']:
                grouped_tables[uat_status].append((os_ticket, script_mapping, uat_status))
            else:
                print(f"Warning: Unexpected UAT status '{uat_status}' encountered.")
    
    # Create tables for each UAT status
    for status, data_list in grouped_tables.items():
        if data_list:
            doc.add_heading(status, level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'OS Ticket Number'
            hdr_cells[1].text = 'Script Mapping'
            hdr_cells[2].text = 'UAT Status'

            # Populate the table with data
            for data_row in data_list:
                row_cells = table.add_row().cells
                row_cells[0].text = str(data_row[0])  # OS Ticket Number
                row_cells[1].text = str(data_row[1])  # Script Mapping
                row_cells[2].text = data_row[2]       # UAT Status

    # Center-align the content in each cell
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
