import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def extract_tables_from_excel(excel_file_path):
    wb = openpyxl.load_workbook(excel_file_path)
    extracted_tables = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        table_data = {}
        for row in ws.iter_rows(values_only=True):
            os_ticket, script_mapping, uat_status = row[0], row[2], row[6]  # Assuming the indexes match the column order
            uat_status = uat_status.lower()  # Convert to lowercase
            if uat_status not in table_data:
                table_data[uat_status] = {'os_tickets': [], 'script_mappings': []}
            table_data[uat_status]['os_tickets'].append(os_ticket)
            table_data[uat_status]['script_mappings'].append(script_mapping)
        extracted_tables.append((sheet, table_data))
    return extracted_tables


def create_word_document(tables, output_path):
    doc = Document()

    for sheet_name, table_data in tables:
        doc.add_heading(sheet_name, level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'OS Ticket Number'
        hdr_cells[1].text = 'Script Mapping'
        hdr_cells[2].text = 'UAT Status'

        for status, data in table_data.items():
            os_tickets = ', '.join(str(ticket) for ticket in data['os_tickets'])
            script_mappings = ', '.join(data['script_mappings'])
            row_cells = table.add_row().cells
            row_cells[0].text = os_tickets
            row_cells[1].text = script_mappings
            row_cells[2].text = status.capitalize()  # Capitalize the UAT Status for consistency

    # Merge rows and format cells
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
